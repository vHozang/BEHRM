from __future__ import annotations

import io
import json
import mimetypes
import os
import re
from dataclasses import asdict, dataclass, field
from html import unescape
from typing import Any, Callable, Dict, List, Optional

import fitz
import requests
from docx import Document


@dataclass
class ParsedDocument:
    text: str
    markdown: str
    content_blocks: List[Dict[str, Any]]
    parser_name: str
    parser_version: str = ""
    used_ocr: bool = False
    quality_score: float = 0.0
    warnings: List[str] = field(default_factory=list)

    def metadata(self) -> Dict[str, Any]:
        data = asdict(self)
        data.pop("text", None)
        data.pop("markdown", None)
        data.pop("content_blocks", None)
        return data


def _normalize_whitespace(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[\t\r\f\v]+", " ", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _quality_score(text: str, page_count: int) -> float:
    if not text.strip():
        return 0.0
    printable = sum(1 for char in text if char.isprintable() or char in "\n\t") / max(len(text), 1)
    alphanumeric = sum(1 for char in text if char.isalnum()) / max(len(text), 1)
    length_factor = min(1.0, len(text.strip()) / max(600, page_count * 300))
    return round(max(0.0, min(1.0, printable * 0.35 + alphanumeric * 0.25 + length_factor * 0.40)), 4)


def _strip_html(value: str) -> str:
    return _normalize_whitespace(unescape(re.sub(r"<[^>]+>", " ", value or "")))


def _block_text(item: Dict[str, Any]) -> str:
    block_type = str(item.get("type", "text"))
    if block_type in {"text", "equation", "header", "footer", "page_number", "aside_text", "page_footnote"}:
        return str(item.get("text") or "")
    if block_type == "table":
        caption = " ".join(item.get("table_caption") or [])
        body = _strip_html(str(item.get("table_body") or ""))
        return _normalize_whitespace(f"{caption}\n{body}")
    if block_type == "list":
        return "\n".join(str(value) for value in item.get("list_items") or [])
    if block_type == "code":
        return str(item.get("code_body") or "")
    if block_type in {"image", "chart"}:
        captions = item.get("image_caption") or item.get("chart_caption") or []
        content = str(item.get("content") or "")
        return _normalize_whitespace("\n".join([*map(str, captions), content]))
    return str(item.get("text") or item.get("content") or "")


def _content_blocks(content_list: Any) -> List[Dict[str, Any]]:
    if isinstance(content_list, str):
        try:
            content_list = json.loads(content_list)
        except json.JSONDecodeError:
            return []
    if not isinstance(content_list, list):
        return []

    blocks: List[Dict[str, Any]] = []
    for item in content_list:
        if not isinstance(item, dict):
            continue
        text = _block_text(item)
        if not text:
            continue
        blocks.append({
            "type": str(item.get("type") or "text"),
            "text": text,
            "page": int(item.get("page_idx", 0)) + 1,
            "bbox": item.get("bbox") if isinstance(item.get("bbox"), list) else None,
            "level": item.get("text_level"),
        })
    return blocks


class MinerUAdapter:
    def __init__(self) -> None:
        self.base_url = os.getenv("MINERU_URL", "").rstrip("/")
        self.backend = os.getenv("MINERU_BACKEND", "pipeline")
        self.timeout = max(30, int(os.getenv("MINERU_TIMEOUT", "600")))

    @property
    def enabled(self) -> bool:
        return bool(self.base_url)

    def health(self) -> Dict[str, Any]:
        if not self.enabled:
            return {"enabled": False, "status": "disabled"}
        try:
            response = requests.get(f"{self.base_url}/health", timeout=5)
            payload = response.json() if response.headers.get("content-type", "").startswith("application/json") else {}
            return {
                "enabled": True,
                "status": "ok" if response.ok else "error",
                "http_status": response.status_code,
                "version": payload.get("version"),
            }
        except Exception as exc:
            return {"enabled": True, "status": "unavailable", "error": str(exc)}

    def parse(self, filename: str, file_bytes: bytes) -> ParsedDocument:
        if not self.enabled:
            raise RuntimeError("MINERU_URL is not configured")

        mime_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
        response = requests.post(
            f"{self.base_url}/file_parse",
            files={"files": (filename, file_bytes, mime_type)},
            data={
                "backend": self.backend,
                "parse_method": "auto",
                "formula_enable": "false",
                "table_enable": "true",
                "image_analysis": "false",
                "return_md": "true",
                "return_content_list": "true",
                "return_middle_json": "false",
                "return_model_output": "false",
                "return_images": "false",
                "response_format_zip": "false",
            },
            timeout=self.timeout,
        )
        if not response.ok:
            detail = response.text[:500] if response.text else f"HTTP {response.status_code}"
            raise RuntimeError(f"MinerU parse failed: {detail}")

        payload = response.json()
        result_map = payload.get("results") or {}
        result = next(iter(result_map.values()), {}) if isinstance(result_map, dict) else {}
        blocks = _content_blocks(result.get("content_list"))
        markdown = str(result.get("md_content") or "")
        text = "\n".join(block["text"] for block in blocks) or markdown
        text = _normalize_whitespace(text)
        page_count = max([int(block.get("page") or 1) for block in blocks], default=1)
        return ParsedDocument(
            text=text,
            markdown=markdown,
            content_blocks=blocks,
            parser_name="mineru",
            parser_version=str(payload.get("version") or ""),
            used_ocr=self.backend == "pipeline",
            quality_score=_quality_score(text, page_count),
        )


class ResumeDocumentParser:
    def __init__(self) -> None:
        self.mineru = MinerUAdapter()
        self.strategy = os.getenv("MINERU_STRATEGY", "auto").strip().lower()
        self.quality_threshold = float(os.getenv("MINERU_QUALITY_THRESHOLD", "0.72"))

    def parse(self, filename: str, file_bytes: bytes, force_mineru: bool = False) -> ParsedDocument:
        fast = self._fast_parse(filename, file_bytes)
        should_use_mineru = force_mineru or self.strategy == "prefer"
        if self.strategy == "auto" and fast.quality_score < self.quality_threshold:
            should_use_mineru = True
        if self.strategy == "disabled":
            should_use_mineru = False

        if should_use_mineru and self.mineru.enabled:
            try:
                parsed = self.mineru.parse(filename, file_bytes)
                if parsed.quality_score >= fast.quality_score or force_mineru:
                    return parsed
                fast.warnings.append("MinerU output had lower text quality; fast parser output was retained.")
            except Exception as exc:
                fast.warnings.append(f"MinerU unavailable; fallback parser used: {exc}")
        elif should_use_mineru and not self.mineru.enabled:
            fast.warnings.append("MinerU was requested but MINERU_URL is not configured.")
        return fast

    def _fast_parse(self, filename: str, file_bytes: bytes) -> ParsedDocument:
        lower = filename.lower()
        if lower.endswith(".pdf"):
            return self._parse_pdf(file_bytes)
        if lower.endswith(".docx"):
            return self._parse_docx(file_bytes)
        raise ValueError(f"Unsupported file type: {filename}")

    def _parse_pdf(self, file_bytes: bytes) -> ParsedDocument:
        blocks: List[Dict[str, Any]] = []
        text_parts: List[str] = []
        with fitz.open(stream=file_bytes, filetype="pdf") as document:
            for page_index, page in enumerate(document):
                page_text = page.get_text("text").strip()
                if page_text:
                    text_parts.append(page_text)
                    blocks.append({"type": "text", "text": page_text, "page": page_index + 1, "bbox": None, "level": None})
            page_count = len(document)
        text = _normalize_whitespace("\n".join(text_parts))
        return ParsedDocument(
            text=text,
            markdown=text,
            content_blocks=blocks,
            parser_name="pymupdf",
            parser_version=fitz.VersionBind,
            quality_score=_quality_score(text, page_count),
            warnings=[] if text else ["No embedded PDF text was found; OCR is recommended."],
        )

    def _parse_docx(self, file_bytes: bytes) -> ParsedDocument:
        document = Document(io.BytesIO(file_bytes))
        blocks = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append({"type": "text", "text": text, "page": 1, "bbox": None, "level": None})
        text = _normalize_whitespace("\n".join(block["text"] for block in blocks))
        return ParsedDocument(
            text=text,
            markdown=text,
            content_blocks=blocks,
            parser_name="python-docx",
            parser_version="",
            quality_score=_quality_score(text, 1),
            warnings=[] if text else ["No DOCX paragraph text was found."],
        )


def redact_scoring_text(text: str, candidate_name: str = "") -> str:
    redacted = re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "[EMAIL]", text)
    redacted = re.sub(r"(?<!\d)(?:\+?84|0)(?:[ .-]?\d){8,10}(?!\d)", "[PHONE]", redacted)
    redacted = re.sub(
        r"(?i)\b(?:date of birth|dob|ngay sinh|giới tính|gioi tinh|gender|marital status|tinh trang hon nhan)\s*[:\-]?[^\n]{0,80}",
        "[PROTECTED_ATTRIBUTE_REMOVED]",
        redacted,
    )
    if candidate_name and len(candidate_name.strip()) >= 4:
        redacted = re.sub(re.escape(candidate_name.strip()), "[CANDIDATE]", redacted, flags=re.IGNORECASE)
    return _normalize_whitespace(redacted)


def build_cv_profile(
    parsed: ParsedDocument,
    filename: str,
    candidate_name: str,
    email: str,
    skills: List[str],
    years_experience: int,
    phone_extractor: Optional[Callable[[str], str]] = None,
) -> Dict[str, Any]:
    blocks = parsed.content_blocks or [{"text": parsed.text, "page": 1, "type": "text"}]

    def evidence_for(term: str) -> Dict[str, Any]:
        normalized_term = term.lower()
        for block in blocks:
            for line in str(block.get("text") or "").splitlines():
                if normalized_term in line.lower():
                    return {"text": line.strip()[:500], "source_page": int(block.get("page") or 1)}
        return {"text": "No supporting evidence found", "source_page": None}

    skill_items = []
    for skill in skills:
        evidence = evidence_for(skill)
        skill_items.append({"name": skill, "evidence": evidence["text"], "source_page": evidence["source_page"]})

    phone = phone_extractor(parsed.text) if phone_extractor else ""
    warnings = list(parsed.warnings)
    if not skills:
        warnings.append("No recognized skills were extracted.")
    if parsed.quality_score < 0.6:
        warnings.append("Document extraction quality is low; human verification is required.")

    completeness = sum(bool(value) for value in [candidate_name, email, skills, years_experience]) / 4
    extraction_confidence = round(min(1.0, parsed.quality_score * 0.7 + completeness * 0.3), 4)
    return {
        "schema_version": "cv-profile.v1",
        "personal_info": {
            "full_name": candidate_name,
            "email": email,
            "phone_number": phone,
        },
        "summary": parsed.text[:800],
        "skills": skill_items,
        "work_experience": [{
            "years_total": years_experience,
            "evidence": evidence_for(str(years_experience))["text"] if years_experience else "No supporting evidence found",
        }],
        "education": [],
        "certifications": [],
        "projects": [],
        "source_filename": filename,
        "extraction_confidence": extraction_confidence,
        "warnings": warnings,
        "parser": parsed.metadata(),
    }

